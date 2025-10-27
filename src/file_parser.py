"""
File parser for resume files (PDF, DOC, DOCX, TXT).
"""

import os
import logging
from pathlib import Path
from typing import Optional, Dict, Any

import pypdf
from docx import Document
from rich.console import Console

from .ats_models import FileParseResult

logger = logging.getLogger(__name__)
console = Console()


class ResumeFileParser:
    """Parser for various resume file formats."""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt'}
    
    def __init__(self):
        self.logger = logger
    
    def parse_file(self, file_path: str) -> FileParseResult:
        """
        Parse a resume file and extract text content.
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            FileParseResult with extracted content and metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            return FileParseResult(
                file_path=str(file_path),
                file_type="unknown",
                content="",
                parse_success=False,
                error_message=f"File not found: {file_path}"
            )
        
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.SUPPORTED_EXTENSIONS:
            return FileParseResult(
                file_path=str(file_path),
                file_type=file_extension,
                content="",
                parse_success=False,
                error_message=f"Unsupported file type: {file_extension}"
            )
        
        try:
            if file_extension == '.pdf':
                return self._parse_pdf(file_path)
            elif file_extension == '.docx':
                return self._parse_docx(file_path)
            elif file_extension == '.doc':
                return self._parse_doc(file_path)
            elif file_extension == '.txt':
                return self._parse_txt(file_path)
            else:
                return FileParseResult(
                    file_path=str(file_path),
                    file_type=file_extension,
                    content="",
                    parse_success=False,
                    error_message=f"Parser not implemented for: {file_extension}"
                )
                
        except Exception as e:
            self.logger.error(f"Error parsing file {file_path}: {e}")
            return FileParseResult(
                file_path=str(file_path),
                file_type=file_extension,
                content="",
                parse_success=False,
                error_message=f"Parse error: {str(e)}"
            )
    
    def _parse_pdf(self, file_path: Path) -> FileParseResult:
        """Parse PDF file using pypdf."""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                
                content_parts = []
                metadata = {
                    'num_pages': len(pdf_reader.pages),
                    'title': None,
                    'author': None,
                    'subject': None,
                    'creator': None
                }
                
                # Extract metadata if available
                if pdf_reader.metadata:
                    metadata.update({
                        'title': pdf_reader.metadata.get('/Title'),
                        'author': pdf_reader.metadata.get('/Author'),
                        'subject': pdf_reader.metadata.get('/Subject'),
                        'creator': pdf_reader.metadata.get('/Creator')
                    })
                
                # Extract text from all pages
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            content_parts.append(page_text)
                    except Exception as e:
                        self.logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                        continue
                
                content = '\n\n'.join(content_parts)
                
                return FileParseResult(
                    file_path=str(file_path),
                    file_type='pdf',
                    content=content,
                    metadata=metadata,
                    parse_success=True
                )
                
        except Exception as e:
            raise Exception(f"PDF parsing failed: {e}")
    
    def _parse_docx(self, file_path: Path) -> FileParseResult:
        """Parse DOCX file using python-docx."""
        try:
            doc = Document(file_path)
            
            content_parts = []
            metadata = {
                'num_paragraphs': len(doc.paragraphs),
                'num_tables': len(doc.tables),
                'title': None,
                'author': None,
                'subject': None
            }
            
            # Extract metadata if available
            if doc.core_properties:
                metadata.update({
                    'title': doc.core_properties.title,
                    'author': doc.core_properties.author,
                    'subject': doc.core_properties.subject
                })
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                table_text_parts = []
                for row in table.rows:
                    row_text_parts = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text_parts.append(cell.text.strip())
                    if row_text_parts:
                        table_text_parts.append(' | '.join(row_text_parts))
                if table_text_parts:
                    content_parts.append('\n'.join(table_text_parts))
            
            content = '\n\n'.join(content_parts)
            
            return FileParseResult(
                file_path=str(file_path),
                file_type='docx',
                content=content,
                metadata=metadata,
                parse_success=True
            )
            
        except Exception as e:
            raise Exception(f"DOCX parsing failed: {e}")
    
    def _parse_doc(self, file_path: Path) -> FileParseResult:
        """Parse DOC file (legacy Word format)."""
        # Note: python-docx doesn't support .doc files
        # This would require additional libraries like python-docx2txt or antiword
        # For now, we'll return an error with a helpful message
        return FileParseResult(
            file_path=str(file_path),
            file_type='doc',
            content="",
            parse_success=False,
            error_message="DOC file parsing not implemented. Please convert to DOCX format."
        )
    
    def _parse_txt(self, file_path: Path) -> FileParseResult:
        """Parse plain text file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                raise Exception("Could not decode file with any supported encoding")
            
            metadata = {
                'file_size': file_path.stat().st_size,
                'encoding': encoding
            }
            
            return FileParseResult(
                file_path=str(file_path),
                file_type='txt',
                content=content,
                metadata=metadata,
                parse_success=True
            )
            
        except Exception as e:
            raise Exception(f"TXT parsing failed: {e}")
    
    def get_supported_extensions(self) -> set:
        """Get list of supported file extensions."""
        return self.SUPPORTED_EXTENSIONS.copy()
    
    def is_supported(self, file_path: str) -> bool:
        """Check if file type is supported."""
        return Path(file_path).suffix.lower() in self.SUPPORTED_EXTENSIONS
